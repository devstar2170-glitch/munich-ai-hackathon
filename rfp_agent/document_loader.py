from pathlib import Path


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md"}


def load_document(file_path: str) -> tuple[str, str]:
    """Load a document and return (text_content, mime_type).

    For PDFs, also returns the raw path so the caller can use the Gemini Files API
    for native PDF understanding (much better than text extraction for complex layouts).
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type '{suffix}'. Supported: {SUPPORTED_EXTENSIONS}")

    if suffix == ".pdf":
        return _load_pdf(path), "application/pdf"
    elif suffix in (".docx", ".doc"):
        return _load_docx(path), "text/plain"
    else:
        return path.read_text(encoding="utf-8"), "text/plain"


def _load_pdf(path: Path) -> str:
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("Install pdfplumber for PDF support: pip install pdfplumber")

    pages = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text and text.strip():
                pages.append(f"[Page {i + 1}]\n{text.strip()}")

    if not pages:
        raise ValueError(f"No text could be extracted from {path}. The PDF may be scanned/image-based.")

    return "\n\n".join(pages)


def _load_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("Install python-docx for DOCX support: pip install python-docx")

    doc = Document(path)
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            # Preserve heading structure
            if para.style.name.startswith("Heading"):
                level = para.style.name.split()[-1] if para.style.name != "Heading" else "1"
                parts.append(f"\n{'#' * int(level)} {para.text.strip()}")
            else:
                parts.append(para.text.strip())

    # Also extract tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            parts.append("\n[TABLE]\n" + "\n".join(rows) + "\n[/TABLE]")

    return "\n\n".join(parts)
